"""
Document processing utilities for RAG-based voice agent.
Handles document loading, text extraction, and chunking.
"""

import logging
from pathlib import Path
from typing import Dict, List, Tuple

import PyPDF2
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Try to import additional PDF libraries for fallback
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document loading, text extraction, and chunking."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file with multiple fallback methods.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        # Method 1: Try PyPDF2 first
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:  # Only add if text was extracted
                        text += page_text + "\n"

                if text.strip():  # If we got some text, return it
                    logger.info(f"Successfully extracted text from PDF using PyPDF2: {file_path}")
                    return text.strip()

        except Exception as e:
            logger.warning(f"PyPDF2 failed for {file_path}: {e}")

        # Method 2: Try PyMuPDF (fitz) if available
        if HAS_PYMUPDF:
            try:
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()

                if text.strip():
                    logger.info(f"Successfully extracted text from PDF using PyMuPDF: {file_path}")
                    return text.strip()

            except Exception as e:
                logger.warning(f"PyMuPDF failed for {file_path}: {e}")

        # Method 3: Try pdfplumber if available
        if HAS_PDFPLUMBER:
            try:
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"

                if text.strip():
                    logger.info(f"Successfully extracted text from PDF using pdfplumber: {file_path}")
                    return text.strip()

            except Exception as e:
                logger.warning(f"pdfplumber failed for {file_path}: {e}")

        # If all methods failed
        logger.error(f"All PDF extraction methods failed for {file_path}")
        return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file including paragraphs and tables.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            doc = Document(file_path)
            texts = []

            # Extract from normal paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # skip empty
                    texts.append(paragraph.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = " ".join([p.text for p in cell.paragraphs if p.text.strip()])
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        texts.append("\t".join(row_text))  # tab-separated like a row

            return "\n".join(texts).strip()

        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return ""

    def extract_text_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Extracted text
        """
        try:
            with open(file_path, encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            return ""

    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from file based on extension.
        
        Args:
            file_path: Path to file
            
        Returns:
            Extracted text
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        logger.info(f"Attempting to extract text from: {file_path} (extension: {extension})")

        if extension == '.pdf':
            result = self.extract_text_from_pdf(str(file_path))
            logger.info(f"PDF text extraction result length: {len(result)} characters")
            return result
        elif extension == '.docx':
            result = self.extract_text_from_docx(str(file_path))
            logger.info(f"DOCX text extraction result length: {len(result)} characters")
            return result
        elif extension == '.txt':
            result = self.extract_text_from_txt(str(file_path))
            logger.info(f"TXT text extraction result length: {len(result)} characters")
            return result
        else:
            logger.warning(f"Unsupported file type: {extension}")
            return ""

    def chunk_text(self, text: str, source: str = "") -> List[Tuple[str, Dict]]:
        """
        Split text into chunks.
        
        Args:
            text: Text to split
            source: Source file name or identifier
            
        Returns:
            List of tuples (chunk_text, metadata)
        """
        if not text.strip():
            return []

        chunks = self.text_splitter.split_text(text)

        # Create metadata for each chunk
        result = []
        for i, chunk in enumerate(chunks):
            metadata = {
                "source": source,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "chunk_size": len(chunk)
            }
            result.append((chunk, metadata))

        return result

    def process_file(self, file_path: str) -> List[Tuple[str, Dict]]:
        """
        Process a file: extract text and create chunks.
        
        Args:
            file_path: Path to file to process
            
        Returns:
            List of tuples (chunk_text, metadata)
        """
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error(f"File does not exist: {file_path}")
            return []

        # Extract text
        text = self.extract_text_from_file(str(file_path))

        if not text:
            logger.warning(f"No text extracted from file: {file_path}")
            return []

        # Create chunks
        chunks = self.chunk_text(text, source=file_path.name)

        logger.info(f"Processed file {file_path.name}: {len(chunks)} chunks created")
        return chunks

    def process_directory(self, directory_path: str, extensions: List[str] = None) -> List[Tuple[str, Dict]]:
        """
        Process all supported files in a directory.
        
        Args:
            directory_path: Path to directory
            extensions: List of file extensions to process (default: ['.pdf', '.docx', '.txt'])
            
        Returns:
            List of tuples (chunk_text, metadata)
        """
        if extensions is None:
            extensions = ['.pdf', '.docx', '.txt']

        directory_path = Path(directory_path)

        if not directory_path.exists():
            logger.error(f"Directory does not exist: {directory_path}")
            return []

        all_chunks = []
        
        logger.info(f"Processing directory: {directory_path} with extensions: {extensions}")

        for file_path in directory_path.rglob('*'):
            if file_path.is_file():
                logger.info(f"Found file: {file_path} (extension: {file_path.suffix.lower()})")
                if file_path.suffix.lower() in extensions:
                    logger.info(f"Processing supported file: {file_path}")
                    chunks = self.process_file(str(file_path))
                    logger.info(f"File {file_path.name} produced {len(chunks)} chunks")
                    all_chunks.extend(chunks)
                else:
                    logger.info(f"Skipping unsupported file: {file_path}")

        logger.info(f"Processed directory {directory_path}: {len(all_chunks)} total chunks")
        return all_chunks
